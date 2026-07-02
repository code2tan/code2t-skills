#!/usr/bin/env python3
"""
WeChat article archive helper.

Capabilities:
- current: archive one public WeChat article as Markdown/HTML/metadata/images, optionally DOCX/PDF/ZIP.
- batch: archive many explicit article links with dedupe, resume, delay, index and failure report.
- history: use user-provided legal cookies/tokens/profile URL to request WeChat history list, then optionally archive recent links.
- album/workflow/browser-context: extract links from album pages or browser-read page exports, then optionally archive.
- publish: create Lark/IMA-ready offline packages and handoff instructions from archive outputs.

This script does not bypass login, paywall, captcha, deleted articles, or platform access control.
"""
from __future__ import annotations

import argparse
import csv
import hashlib
import html
import json
import random
import re
import sys
import time
import zipfile
from dataclasses import asdict, dataclass
from pathlib import Path
from typing import Any, Iterable, Optional
from urllib.parse import parse_qs, urlencode, urlparse, urlunparse

import requests
from bs4 import BeautifulSoup

try:
    from markdownify import markdownify as html_to_markdown
except Exception:  # pragma: no cover
    html_to_markdown = None

try:
    from docx import Document
except Exception:  # pragma: no cover
    Document = None

try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
except Exception:  # pragma: no cover
    SimpleDocTemplate = None


DEFAULT_HEADERS = {
    "User-Agent": (
        "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) "
        "AppleWebKit/537.36 (KHTML, like Gecko) "
        "Chrome/122.0.0.0 Safari/537.36"
    ),
    "Referer": "https://mp.weixin.qq.com/",
    "Accept-Language": "zh-CN,zh;q=0.9,en;q=0.8",
}


@dataclass
class ArticleMeta:
    title: str
    account_name: str
    author: str
    source_url: str
    final_url: str
    publish_time: str
    biz: str
    digest: str


@dataclass
class ArchiveResult:
    url: str
    status: str
    article_dir: str = ""
    title: str = ""
    account_name: str = ""
    publish_time: str = ""
    error: str = ""


def safe_name(value: str, fallback: str = "untitled") -> str:
    value = re.sub(r"[\\/:*?\"<>|\n\r\t]+", "_", value or "").strip(" ._")
    value = re.sub(r"\s+", " ", value)
    return value[:120] or fallback


def make_session(cookie: str = "") -> requests.Session:
    session = requests.Session()
    session.headers.update(DEFAULT_HEADERS)
    if cookie:
        session.headers.update({"Cookie": cookie})
    return session


def request_with_retry(
    session: requests.Session,
    url: str,
    *,
    timeout: int = 30,
    retries: int = 2,
    backoff: float = 1.5,
    allow_redirects: bool = True,
) -> requests.Response:
    last_exc: Optional[Exception] = None
    for attempt in range(retries + 1):
        try:
            response = session.get(url, timeout=timeout, allow_redirects=allow_redirects)
            response.raise_for_status()
            return response
        except Exception as exc:  # noqa: PERF203
            last_exc = exc
            if attempt < retries:
                time.sleep(backoff * (attempt + 1))
    raise RuntimeError(f"请求失败: {last_exc}")


def normalize_url(url: str) -> str:
    """Normalize URL for catalog dedup while preserving WeChat article identity params."""
    url = (url or "").strip()
    if not url:
        return ""
    parsed = urlparse(url)
    query = urlencode(sorted(parse_qs(parsed.query, keep_blank_values=True).items()), doseq=True)
    return urlunparse((parsed.scheme, parsed.netloc, parsed.path, parsed.params, query, ""))


def fetch_html(url: str, *, cookie: str = "", retries: int = 2, timeout: int = 30) -> tuple[str, str]:
    session = make_session(cookie)
    response = request_with_retry(session, url, timeout=timeout, retries=retries)
    response.encoding = response.apparent_encoding or "utf-8"
    return response.text, response.url


def text_of(soup: BeautifulSoup, selector: str) -> str:
    node = soup.select_one(selector)
    return node.get_text(" ", strip=True) if node else ""


def regex_first(pattern: str, text: str) -> str:
    match = re.search(pattern, text, re.S)
    if not match:
        return ""
    return html.unescape(match.group(1)).strip().strip('"')


def extract_biz(final_url: str, page_html: str) -> str:
    query_match = re.search(r"(?:\?|&)__biz=([^&]+)", final_url)
    if query_match:
        return query_match.group(1)
    for pattern in [
        r"var\s+biz\s*=\s*['\"]([^'\"]+)",
        r"__biz=([^&'\"]+)",
        r"biz\s*:\s*['\"]([^'\"]+)",
    ]:
        value = regex_first(pattern, page_html)
        if value:
            return value
    return ""


def extract_publish_time(page_html: str) -> str:
    patterns = [
        r"var\s+ct\s*=\s*['\"]?(\d{10})['\"]?",
        r"publish_time\s*=\s*['\"]([^'\"]+)",
        r"oriCreateTime\s*=\s*['\"]?(\d{10})['\"]?",
    ]
    for pattern in patterns:
        value = regex_first(pattern, page_html)
        if not value:
            continue
        if value.isdigit() and len(value) == 10:
            return time.strftime("%Y-%m-%d %H:%M:%S", time.localtime(int(value)))
        return value
    return ""


def extract_digest(page_html: str) -> str:
    for pattern in [
        r"var\s+msg_desc\s*=\s*['\"](.*?)['\"]\s*;",
        r"meta\s+name=\"description\"\s+content=\"(.*?)\"",
    ]:
        value = regex_first(pattern, page_html)
        if value:
            return value
    return ""


def extract_meta(soup: BeautifulSoup, page_html: str, source_url: str, final_url: str) -> ArticleMeta:
    title = text_of(soup, "#activity-name") or regex_first(r"var\s+msg_title\s*=\s*['\"](.*?)['\"]\s*;", page_html)
    account_name = text_of(soup, "#js_name") or regex_first(r"var\s+nickname\s*=\s*['\"](.*?)['\"]\s*;", page_html)
    author = text_of(soup, "#js_author_name") or regex_first(r"var\s+user_name\s*=\s*['\"](.*?)['\"]\s*;", page_html)
    return ArticleMeta(
        title=title or "untitled",
        account_name=account_name or "unknown-account",
        author=author or "",
        source_url=source_url,
        final_url=final_url,
        publish_time=extract_publish_time(page_html),
        biz=extract_biz(final_url, page_html),
        digest=extract_digest(page_html),
    )


def image_extension(content_type: str, url: str) -> str:
    ct = (content_type or "").split(";")[0].lower()
    mapping = {
        "image/jpeg": ".jpg",
        "image/png": ".png",
        "image/gif": ".gif",
        "image/webp": ".webp",
        "image/svg+xml": ".svg",
    }
    if ct in mapping:
        return mapping[ct]
    suffix = Path(urlparse(url).path).suffix.lower()
    return suffix if suffix in {".jpg", ".jpeg", ".png", ".gif", ".webp", ".svg"} else ".jpg"


def iter_images(content: BeautifulSoup) -> Iterable[Any]:
    yield from content.select("img")


def download_images(
    content: BeautifulSoup,
    article_dir: Path,
    *,
    dry_run: bool = False,
    retries: int = 2,
    cookie: str = "",
) -> list[dict[str, str]]:
    image_dir = article_dir / "images"
    image_dir.mkdir(parents=True, exist_ok=True)
    session = make_session(cookie)
    results: list[dict[str, str]] = []
    seen: dict[str, str] = {}
    for index, img in enumerate(iter_images(content), start=1):
        url = img.get("data-src") or img.get("src") or img.get("data-original")
        if not url or url.startswith("data:"):
            continue
        if url.startswith("//"):
            url = "https:" + url
        if url in seen:
            img["src"] = seen[url]
            continue
        local_rel = f"images/{index:03d}.jpg"
        status = "pending" if dry_run else "ok"
        error = ""
        if not dry_run:
            try:
                resp = request_with_retry(session, url, retries=retries)
                ext = image_extension(resp.headers.get("Content-Type", ""), url)
                local_rel = f"images/{index:03d}{ext}"
                (article_dir / local_rel).write_bytes(resp.content)
            except Exception as exc:
                status = "failed"
                error = str(exc)
        seen[url] = local_rel
        img["src"] = local_rel
        for attr in ["data-src", "data-original", "data-ratio", "data-w"]:
            img.attrs.pop(attr, None)
        results.append({"source": url, "local": local_rel, "status": status, "error": error})
    return results


def cleanup_content(content: BeautifulSoup) -> list[dict[str, str]]:
    degraded: list[dict[str, str]] = []
    for tag in content.select("script, style"):
        tag.decompose()
    for node in content.select("mpvoice, mp-video, iframe, qqmusic, wx-open-launch-weapp"):
        degraded.append({"tag": node.name, "note": "unsupported embedded content replaced by placeholder"})
        placeholder = content.new_tag("p")
        placeholder.string = f"[Unsupported embedded content: {node.name}]"
        node.replace_with(placeholder)
    return degraded


def normalize_wechat_code_blocks(content_html: str) -> str:
    """WeChat renders multi-line code as adjacent <code> siblings inside <pre>
    (<pre><code>line1</code><code>line2</code></pre>) with no separator between them.
    markdownify concatenates them into a single line. Restore a newline between
    each </code> and the next <code> so code blocks keep their line breaks."""
    def _fix_pre(match: re.Match) -> str:
        pre = match.group(0)
        return re.sub(r"</code>(\s*)<code", r"</code>\n<code", pre)

    return re.sub(r"<pre[^>]*>.*?</pre>", _fix_pre, content_html, flags=re.S)


def build_markdown(meta: ArticleMeta, content_html: str) -> str:
    if html_to_markdown:
        body = html_to_markdown(normalize_wechat_code_blocks(content_html), heading_style="ATX")
        # WeChat components (e.g. 阅读推荐) emit runs of empty <pre> blocks that
        # markdownify turns into adjacent ``` fences with no code between them.
        # Collapse any chain of fences that contains only whitespace, leaving
        # real code blocks (fences with non-whitespace between) untouched.
        body = re.sub(r"(?:```[^\S\r\n]*(?:\r?\n)*)+```", "", body)
    else:
        soup = BeautifulSoup(content_html, "html.parser")
        body = soup.get_text("\n", strip=True)
    frontmatter = {
        "title": meta.title,
        "account_name": meta.account_name,
        "author": meta.author,
        "source_url": meta.source_url,
        "final_url": meta.final_url,
        "publish_time": meta.publish_time,
        "biz": meta.biz,
        "digest": meta.digest,
    }
    lines = ["---"]
    for key, value in frontmatter.items():
        safe_value = str(value).replace('"', '\\"')
        lines.append(f'{key}: "{safe_value}"')
    lines.extend(["---", "", body.strip(), ""])
    return "\n".join(lines)


def write_docx(article_dir: Path, meta: ArticleMeta, markdown: str) -> str:
    if Document is None:
        return "python-docx unavailable"
    doc = Document()
    doc.add_heading(meta.title, level=1)
    doc.add_paragraph(f"公众号：{meta.account_name}")
    if meta.author:
        doc.add_paragraph(f"作者：{meta.author}")
    if meta.publish_time:
        doc.add_paragraph(f"发布时间：{meta.publish_time}")
    doc.add_paragraph(f"原文：{meta.source_url}")
    doc.add_paragraph("")
    body = re.sub(r"^---.*?---\s*", "", markdown, flags=re.S).strip()
    for para in body.split("\n\n"):
        text = re.sub(r"!\[[^\]]*\]\([^)]*\)", "[图片见 images 目录]", para).strip()
        if text:
            doc.add_paragraph(text)
    doc.save(article_dir / "article.docx")
    return "ok"


def write_pdf(article_dir: Path, meta: ArticleMeta, markdown: str) -> str:
    if SimpleDocTemplate is None:
        return "reportlab unavailable"
    pdf_path = article_dir / "article.pdf"
    doc = SimpleDocTemplate(str(pdf_path), pagesize=A4)
    styles = getSampleStyleSheet()
    story: list[Any] = []
    story.append(Paragraph(html.escape(meta.title), styles["Title"]))
    story.append(Spacer(1, 12))
    info = f"公众号：{meta.account_name}<br/>发布时间：{meta.publish_time}<br/>原文：{meta.source_url}"
    story.append(Paragraph(info, styles["Normal"]))
    story.append(Spacer(1, 12))
    body = re.sub(r"^---.*?---\s*", "", markdown, flags=re.S).strip()
    body = re.sub(r"!\[[^\]]*\]\([^)]*\)", "[图片见 images 目录]", body)
    for para in body.split("\n\n")[:300]:
        text = html.escape(para.strip()).replace("\n", "<br/>")
        if text:
            story.append(Paragraph(text, styles["BodyText"]))
            story.append(Spacer(1, 8))
    doc.build(story)
    return "ok"


def zip_directory(src_dir: Path, zip_path: Optional[Path] = None) -> Path:
    # Do not use Path.with_suffix here: article titles may contain dots such as "GLM-5.2",
    # and with_suffix would truncate the directory name before the last dot.
    zip_path = zip_path or src_dir.parent / f"{src_dir.name}.zip"
    if zip_path.exists():
        zip_path.unlink()
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
        for file in src_dir.rglob("*"):
            if file.is_file() and file.name != zip_path.name:
                zf.write(file, file.relative_to(src_dir.parent))
    return zip_path


def archive_article(
    url: str,
    output_dir: Path,
    *,
    dry_run: bool = False,
    retries: int = 2,
    cookie: str = "",
    formats: Optional[set[str]] = None,
    make_zip: bool = False,
) -> tuple[Path, ArticleMeta]:
    formats = formats or {"md", "html", "json"}
    page_html, final_url = fetch_html(url, cookie=cookie, retries=retries)
    soup = BeautifulSoup(page_html, "html.parser")
    content = soup.select_one("#js_content")
    if content is None:
        raise RuntimeError("未找到公众号正文节点 #js_content；可能需要登录、文章不可访问或页面结构已变化。")
    meta = extract_meta(soup, page_html, url, final_url)
    digest = hashlib.sha1(final_url.encode("utf-8")).hexdigest()[:8]
    article_dir = output_dir / safe_name(meta.account_name, "unknown-account") / f"{safe_name(meta.title)}-{digest}"
    article_dir.mkdir(parents=True, exist_ok=True)

    degraded = cleanup_content(content)
    images = download_images(content, article_dir, dry_run=dry_run, retries=retries, cookie=cookie)
    content_html = str(content)
    markdown = build_markdown(meta, content_html)

    if "html" in formats:
        (article_dir / "article.html").write_text(content_html, encoding="utf-8")
    if "md" in formats:
        (article_dir / "article.md").write_text(markdown, encoding="utf-8")
    if "json" in formats:
        (article_dir / "metadata.json").write_text(
            json.dumps({"meta": asdict(meta), "images": images, "degraded": degraded}, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )
    if "docx" in formats:
        docx_status = write_docx(article_dir, meta, markdown)
        if docx_status != "ok":
            (article_dir / "docx-warning.txt").write_text(docx_status, encoding="utf-8")
    if "pdf" in formats:
        pdf_status = write_pdf(article_dir, meta, markdown)
        if pdf_status != "ok":
            (article_dir / "pdf-warning.txt").write_text(pdf_status, encoding="utf-8")
    if make_zip:
        zip_directory(article_dir)
    return article_dir, meta


def read_links(path: Path) -> list[str]:
    links: list[str] = []
    for line in path.read_text(encoding="utf-8").splitlines():
        line = line.strip()
        if not line or line.startswith("#"):
            continue
        links.append(line)
    deduped = list(dict.fromkeys(links))
    return deduped


def load_done(manifest_path: Path) -> set[str]:
    if not manifest_path.exists():
        return set()
    done: set[str] = set()
    with manifest_path.open("r", encoding="utf-8", newline="") as f:
        reader = csv.DictReader(f)
        for row in reader:
            if row.get("status") == "ok" and row.get("url"):
                done.add(row["url"])
    return done


def append_manifest(manifest_path: Path, result: ArchiveResult) -> None:
    exists = manifest_path.exists()
    with manifest_path.open("a", encoding="utf-8", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=list(asdict(result).keys()))
        if not exists:
            writer.writeheader()
        writer.writerow(asdict(result))


def stable_article_key(url: str = "", title: str = "", account_name: str = "", publish_time: str = "") -> str:
    """Create a stable dedup key without storing cookies or private tokens."""
    import hashlib

    base = "|".join([normalize_url(url), title.strip(), account_name.strip(), publish_time.strip()])
    return hashlib.sha256(base.encode("utf-8")).hexdigest()[:16]


def load_catalog(catalog_path: Path) -> dict[str, dict[str, Any]]:
    if not catalog_path.exists():
        return {}
    try:
        data = json.loads(catalog_path.read_text(encoding="utf-8"))
    except Exception:
        return {}
    if isinstance(data, dict) and isinstance(data.get("items"), dict):
        return data["items"]
    return {}


def write_catalog(catalog_path: Path, catalog: dict[str, dict[str, Any]]) -> None:
    payload = {
        "version": 1,
        "count": len(catalog),
        "items": catalog,
    }
    catalog_path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")


def catalog_has_url(catalog: dict[str, dict[str, Any]], url: str) -> bool:
    normalized = normalize_url(url)
    return any(item.get("url") == normalized or item.get("source_url") == normalized for item in catalog.values())


def write_incremental_report(batch_dir: Path, rows: list[dict[str, Any]]) -> None:
    lines = ["# 增量归档报告", "", "| 状态 | 标题 | 公众号 | 原文 | 说明 |", "|---|---|---|---|---|"]
    for row in rows:
        lines.append(
            f"| {row.get('status','')} | {row.get('title','')} | {row.get('account_name','')} | {row.get('url','')} | {row.get('note','')} |"
        )
    (batch_dir / "incremental-report.md").write_text("\n".join(lines) + "\n", encoding="utf-8")
    (batch_dir / "incremental-report.json").write_text(json.dumps(rows, ensure_ascii=False, indent=2), encoding="utf-8")


def write_index(batch_dir: Path, results: list[ArchiveResult]) -> None:
    lines = ["# 微信公众号文章归档索引", "", "| 状态 | 标题 | 公众号 | 发布时间 | 本地路径 | 原文 | 错误 |", "|---|---|---|---|---|---|---|"]
    for item in results:
        local = item.article_dir.replace("|", " ")
        lines.append(
            f"| {item.status} | {item.title} | {item.account_name} | {item.publish_time} | {local} | {item.url} | {item.error} |"
        )
    (batch_dir / "index.md").write_text("\n".join(lines) + "\n", encoding="utf-8")


def batch_archive(args: argparse.Namespace) -> int:
    links = read_links(Path(args.links))
    batch_dir = Path(args.output)
    batch_dir.mkdir(parents=True, exist_ok=True)
    manifest_path = batch_dir / "manifest.csv"
    failure_path = batch_dir / "failures.json"
    catalog_path = Path(args.catalog) if getattr(args, "catalog", "") else batch_dir / "catalog.json"
    done = load_done(manifest_path) if args.resume else set()
    catalog = load_catalog(catalog_path) if getattr(args, "incremental", False) else {}
    results: list[ArchiveResult] = []
    failures: list[dict[str, str]] = []
    incremental_rows: list[dict[str, Any]] = []
    formats = set(args.formats.split(","))

    for idx, url in enumerate(links, start=1):
        normalized_url = normalize_url(url)
        if getattr(args, "incremental", False) and catalog_has_url(catalog, normalized_url):
            result = ArchiveResult(url=url, status="skipped", error="incremental: already in catalog")
            results.append(result)
            incremental_rows.append({"status": "skipped", "url": normalized_url, "note": "already in catalog"})
            continue
        if url in done:
            result = ArchiveResult(url=url, status="skipped", error="resume: already archived")
            results.append(result)
            incremental_rows.append({"status": "skipped", "url": normalized_url, "note": "resume manifest hit"})
            continue
        if "mp.weixin.qq.com" not in url:
            result = ArchiveResult(url=url, status="failed", error="unsupported url")
            append_manifest(manifest_path, result)
            results.append(result)
            failures.append(asdict(result))
            continue
        try:
            article_dir, meta = archive_article(
                url,
                batch_dir / "articles",
                dry_run=args.dry_run,
                retries=args.retries,
                cookie=args.cookie or "",
                formats=formats,
                make_zip=args.article_zip,
            )
            result = ArchiveResult(
                url=url,
                status="ok",
                article_dir=str(article_dir),
                title=meta.title,
                account_name=meta.account_name,
                publish_time=meta.publish_time,
            )
            if getattr(args, "incremental", False):
                key = stable_article_key(url=meta.final_url or url, title=meta.title, account_name=meta.account_name, publish_time=meta.publish_time)
                catalog[key] = {
                    "url": normalized_url,
                    "source_url": normalize_url(meta.source_url or url),
                    "final_url": normalize_url(meta.final_url or url),
                    "title": meta.title,
                    "account_name": meta.account_name,
                    "publish_time": meta.publish_time,
                    "article_dir": str(article_dir),
                }
                incremental_rows.append({"status": "added", "url": normalized_url, "title": meta.title, "account_name": meta.account_name, "note": key})
        except Exception as exc:
            result = ArchiveResult(url=url, status="failed", error=str(exc))
            failures.append(asdict(result))
        append_manifest(manifest_path, result)
        results.append(result)
        if idx < len(links) and args.delay_max > 0:
            time.sleep(random.uniform(args.delay_min, args.delay_max))

    failure_path.write_text(json.dumps(failures, ensure_ascii=False, indent=2), encoding="utf-8")
    if getattr(args, "incremental", False):
        write_catalog(catalog_path, catalog)
        write_incremental_report(batch_dir, incremental_rows)
    write_index(batch_dir, results)
    if args.zip:
        zip_directory(batch_dir)
    print(str(batch_dir))
    return 0 if not failures else 1


def extract_album_items(page_html: str) -> list[dict[str, str]]:
    """Extract article links from a public WeChat album page HTML.

    WeChat album pages often embed article links in escaped JavaScript strings. This parser
    intentionally relies on generic URL/title patterns rather than private APIs.
    """
    variants = [page_html, html.unescape(page_html), page_html.replace("\\/", "/"), html.unescape(page_html).replace("\\/", "/")]
    patterns = [
        r"https?://mp\.weixin\.qq\.com/s\?[^\"'<>\\]+",
        r"https?://mp\.weixin\.qq\.com/s/[^\"'<>\\]+",
        r"//mp\.weixin\.qq\.com/s\?[^\"'<>\\]+",
        r"//mp\.weixin\.qq\.com/s/[^\"'<>\\]+",
    ]
    links: list[str] = []
    for body in variants:
        for pattern in patterns:
            for link in re.findall(pattern, body):
                link = html.unescape(link).replace("\\/", "/")
                if link.startswith("//"):
                    link = "https:" + link
                link = link.split("&amp;")[0] if "&amp;" in link else link
                links.append(link)
    deduped_links = list(dict.fromkeys(links))

    # Best-effort title extraction near content_url fields. If unavailable, leave title empty.
    title_by_link: dict[str, str] = {}
    normalized = html.unescape(page_html).replace("\\/", "/")
    for link in deduped_links:
        pos = normalized.find(link)
        if pos >= 0:
            window = normalized[max(0, pos - 1200): pos + 1200]
            candidates = re.findall(r"[\"']title[\"']\s*[:=]\s*[\"']([^\"']{1,160})[\"']", window)
            if candidates:
                title_by_link[link] = html.unescape(candidates[-1])
    return [{"title": title_by_link.get(link, ""), "link": link} for link in deduped_links]


def fetch_album(args: argparse.Namespace) -> int:
    output = Path(args.output)
    output.mkdir(parents=True, exist_ok=True)
    try:
        page_html, final_url = fetch_html(args.url, cookie=args.cookie or "", retries=args.retries)
    except Exception as exc:
        (output / "album-error.txt").write_text(str(exc), encoding="utf-8")
        print(f"ERROR: 合集页请求失败: {exc}", file=sys.stderr)
        return 1
    items = extract_album_items(page_html)
    if args.limit > 0:
        items = items[: args.limit]
    links = [item["link"] for item in items]
    (output / "album.html").write_text(page_html, encoding="utf-8")
    (output / "album-links.txt").write_text("\n".join(links) + ("\n" if links else ""), encoding="utf-8")
    (output / "album-items.json").write_text(
        json.dumps({"source_url": args.url, "final_url": final_url, "count": len(items), "items": items}, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    if args.archive and links:
        batch_args = argparse.Namespace(
            links=str(output / "album-links.txt"),
            output=str(output / "archive"),
            resume=True,
            dry_run=args.dry_run,
            retries=args.retries,
            cookie=args.cookie or "",
            formats=args.formats,
            article_zip=args.article_zip,
            zip=args.zip,
            delay_min=args.delay_min,
            delay_max=args.delay_max,
            incremental=getattr(args, "incremental", False),
            catalog=getattr(args, "catalog", ""),
        )
        return batch_archive(batch_args)
    print(str(output / "album-links.txt"))
    return 0 if links else 1


def workflow_archive(args: argparse.Namespace) -> int:
    """Run an end-to-end workflow: album -> links -> batch -> index -> zip."""
    output = Path(args.output)
    output.mkdir(parents=True, exist_ok=True)
    album_dir = output / "album"
    archive_dir = output / "archive"
    album_args = argparse.Namespace(
        url=args.url,
        output=str(album_dir),
        cookie=args.cookie or "",
        limit=args.limit,
        retries=args.retries,
        archive=False,
        formats=args.formats,
        dry_run=args.dry_run,
        article_zip=args.article_zip,
        zip=False,
        delay_min=args.delay_min,
        delay_max=args.delay_max,
    )
    album_code = fetch_album(album_args)
    links_path = album_dir / "album-links.txt"
    if album_code != 0 or not links_path.exists() or not links_path.read_text(encoding="utf-8").strip():
        (output / "workflow-error.txt").write_text("album step did not produce links\n", encoding="utf-8")
        return 1
    batch_args = argparse.Namespace(
        links=str(links_path),
        output=str(archive_dir),
        resume=True,
        dry_run=args.dry_run,
        retries=args.retries,
        cookie=args.cookie or "",
        formats=args.formats,
        article_zip=args.article_zip,
        zip=args.zip,
        delay_min=args.delay_min,
        delay_max=args.delay_max,
    )
    batch_code = batch_archive(batch_args)
    workflow_manifest = {
        "source_url": args.url,
        "album_dir": str(album_dir),
        "links_path": str(links_path),
        "archive_dir": str(archive_dir),
        "batch_exitcode": batch_code,
        "zip_path": str(archive_dir.parent / f"{archive_dir.name}.zip") if args.zip else "",
    }
    (output / "workflow-manifest.json").write_text(json.dumps(workflow_manifest, ensure_ascii=False, indent=2), encoding="utf-8")
    print(str(output))
    return batch_code


def read_context_payload(path: Path) -> dict[str, str]:
    text = path.read_text(encoding="utf-8", errors="ignore")
    if path.suffix.lower() == ".json":
        try:
            data = json.loads(text)
            return {
                "url": str(data.get("url") or data.get("final_url") or data.get("source_url") or ""),
                "html": str(data.get("html") or data.get("content") or data.get("page_html") or ""),
                "text": str(data.get("text") or data.get("markdown") or data.get("body") or ""),
            }
        except Exception:
            pass
    return {"url": "", "html": text, "text": text}


def browser_context_extract(args: argparse.Namespace) -> int:
    """Extract WeChat links from a saved browser context export.

    The agent should create this export using browser_read_current_page/browser_read_page and save only URL,
    visible text, or page HTML. This function intentionally does not read cookies or browser storage.
    """
    output = Path(args.output)
    output.mkdir(parents=True, exist_ok=True)
    payload = read_context_payload(Path(args.input))
    combined = "\n".join([payload.get("url", ""), payload.get("html", ""), payload.get("text", "")])
    items = extract_album_items(combined)
    # Also keep direct single article URL if current page itself is an article.
    page_url = payload.get("url", "")
    if "mp.weixin.qq.com/s/" in page_url or "mp.weixin.qq.com/s?" in page_url:
        if page_url not in [item["link"] for item in items]:
            items.insert(0, {"title": "current browser page", "link": page_url})
    if args.limit > 0:
        items = items[: args.limit]
    links = [item["link"] for item in items]
    (output / "browser-context-items.json").write_text(
        json.dumps({"source_input": args.input, "count": len(items), "items": items}, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    links_path = output / "browser-context-links.txt"
    links_path.write_text("\n".join(links) + ("\n" if links else ""), encoding="utf-8")
    if args.archive and links:
        batch_args = argparse.Namespace(
            links=str(links_path),
            output=str(output / "archive"),
            resume=True,
            dry_run=args.dry_run,
            retries=args.retries,
            cookie="",
            formats=args.formats,
            article_zip=args.article_zip,
            zip=args.zip,
            delay_min=args.delay_min,
            delay_max=args.delay_max,
            incremental=getattr(args, "incremental", False),
            catalog=getattr(args, "catalog", ""),
        )
        return batch_archive(batch_args)
    print(str(links_path))
    return 0 if links else 1


def collect_archive_documents(archive_dir: Path) -> list[dict[str, Any]]:
    docs: list[dict[str, Any]] = []
    for md_path in archive_dir.rglob("article.md"):
        article_dir = md_path.parent
        metadata_path = article_dir / "metadata.json"
        meta: dict[str, Any] = {}
        if metadata_path.exists():
            try:
                meta = json.loads(metadata_path.read_text(encoding="utf-8")).get("meta", {})
            except Exception:
                meta = {}
        images = [str(p.relative_to(archive_dir)) for p in (article_dir / "images").glob("*") if p.is_file()] if (article_dir / "images").exists() else []
        docs.append(
            {
                "title": meta.get("title") or article_dir.name,
                "account_name": meta.get("account_name", ""),
                "publish_time": meta.get("publish_time", ""),
                "source_url": meta.get("source_url", ""),
                "markdown_path": str(md_path.relative_to(archive_dir)),
                "metadata_path": str(metadata_path.relative_to(archive_dir)) if metadata_path.exists() else "",
                "images": images,
            }
        )
    return docs


def publish_archive(args: argparse.Namespace) -> int:
    """Create Lark/IMA-ready offline package and handoff instructions from an archive directory."""
    archive_dir = Path(args.archive)
    output = Path(args.output)
    output.mkdir(parents=True, exist_ok=True)
    if not archive_dir.exists():
        print(f"ERROR: archive directory not found: {archive_dir}", file=sys.stderr)
        return 2
    docs = collect_archive_documents(archive_dir)
    (output / "documents.jsonl").write_text("\n".join(json.dumps(doc, ensure_ascii=False) for doc in docs) + ("\n" if docs else ""), encoding="utf-8")
    assets = []
    for doc in docs:
        for image in doc.get("images", []):
            assets.append({"document": doc["markdown_path"], "asset_path": image})
    (output / "assets.jsonl").write_text("\n".join(json.dumps(item, ensure_ascii=False) for item in assets) + ("\n" if assets else ""), encoding="utf-8")
    index_lines = ["# 微信公众号资料库导入索引", "", "| 标题 | 公众号 | 发布时间 | 原文 | Markdown |", "|---|---|---|---|---|"]
    for doc in docs:
        index_lines.append(f"| {doc['title']} | {doc['account_name']} | {doc['publish_time']} | {doc['source_url']} | {doc['markdown_path']} |")
    (output / "index.md").write_text("\n".join(index_lines) + "\n", encoding="utf-8")
    manifest = {
        "target": args.target,
        "archive_dir": str(archive_dir),
        "document_count": len(docs),
        "asset_count": len(assets),
        "mode": args.mode,
        "destination": args.destination,
        "note": "Online write requires user-authorized lark-doc or ima-skill execution by the agent. This package contains no cookies or tokens.",
    }
    (output / "publish-manifest.json").write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")
    handoff_tasks = {
        "target": args.target,
        "mode": args.mode,
        "destination": args.destination,
        "documents": docs,
        "assets": assets,
        "recommended_steps": [
            "Create or open the target Lark document / IMA knowledge base with user authorization.",
            "Write index.md first as the collection landing page.",
            "Import each article.md and upload images according to assets.jsonl.",
            "Use source_url and title as dedup keys before writing.",
        ],
    }
    (output / "handoff-tasks.json").write_text(json.dumps(handoff_tasks, ensure_ascii=False, indent=2), encoding="utf-8")
    handoff = [
        "# 发布交接说明",
        "",
        f"目标：{args.target}",
        f"模式：{args.mode}",
        f"目标位置：{args.destination or '未指定，执行时由用户选择'}",
        f"源归档目录：`{archive_dir}`",
        f"文档数量：{len(docs)}",
        "",
        "## 飞书写入建议",
        "使用 lark-doc 能力创建或追加文档：先写入 index.md，再逐篇写入 Markdown 正文；图片按本地路径作为附件或素材上传。若无权限，保留本离线包给用户手动导入。",
        "",
        "## IMA 知识库写入建议",
        "使用 ima-skill 创建/选择知识库后，将 documents.jsonl 中的 Markdown 逐篇上传；assets.jsonl 作为图片附件映射。若无 API token，交付本离线包。",
        "",
        "## 文件",
        "- index.md",
        "- documents.jsonl",
        "- assets.jsonl",
        "- publish-manifest.json",
        "- handoff-tasks.json",
    ]
    handoff_text = "\n".join(handoff) + "\n"
    (output / "HANDOFF.md").write_text(handoff_text, encoding="utf-8")
    # Backward-compatible lowercase alias for existing automation.
    (output / "handoff.md").write_text(handoff_text, encoding="utf-8")
    if args.zip:
        zip_directory(output)
    print(str(output))
    return 0


def profile_url_with_params(profile_url: str, offset: int, count: int) -> str:
    parsed = urlparse(profile_url)
    query = parse_qs(parsed.query)
    query.update({"action": ["getmsg"], "f": ["json"], "offset": [str(offset)], "count": [str(count)]})
    return urlunparse(parsed._replace(query=urlencode(query, doseq=True)))


def history_url_from_parts(args: argparse.Namespace, offset: int, count: int) -> str:
    if args.profile_url:
        return profile_url_with_params(args.profile_url, offset, count)
    if not args.biz:
        raise RuntimeError("history 模式需要 --profile-url 或 --biz。")
    params = {
        "action": "getmsg",
        "__biz": args.biz,
        "f": "json",
        "offset": offset,
        "count": count,
    }
    optional = {
        "appmsg_token": args.appmsg_token,
        "pass_ticket": args.pass_ticket,
        "uin": args.uin,
        "key": args.key,
        "scene": args.scene,
    }
    params.update({k: v for k, v in optional.items() if v})
    return "https://mp.weixin.qq.com/mp/profile_ext?" + urlencode(params)


def parse_history_items(payload: dict[str, Any]) -> tuple[list[dict[str, Any]], bool]:
    raw = payload.get("general_msg_list") or payload.get("msg_list") or ""
    if isinstance(raw, str):
        try:
            data = json.loads(raw)
        except Exception:
            data = {}
    elif isinstance(raw, dict):
        data = raw
    else:
        data = {}
    messages = data.get("list", []) if isinstance(data, dict) else []
    articles: list[dict[str, Any]] = []
    for message in messages:
        app_msg = message.get("app_msg_ext_info") or {}
        candidates = [app_msg] + list(app_msg.get("multi_app_msg_item_list") or [])
        for item in candidates:
            link = item.get("content_url") or item.get("source_url") or ""
            title = item.get("title") or ""
            if link:
                link = html.unescape(link).replace("\\/", "/")
                articles.append(
                    {
                        "title": title,
                        "link": link,
                        "digest": item.get("digest", ""),
                        "cover": item.get("cover", ""),
                        "datetime": message.get("comm_msg_info", {}).get("datetime", ""),
                    }
                )
    can_msg_continue = bool(payload.get("can_msg_continue", 0))
    return articles, can_msg_continue


def fetch_history(args: argparse.Namespace) -> int:
    if not args.cookie:
        print("ERROR: history 模式需要用户提供合法 Cookie；不会绕过登录态。", file=sys.stderr)
        return 2
    session = make_session(args.cookie)
    output = Path(args.output)
    output.mkdir(parents=True, exist_ok=True)
    all_items: list[dict[str, Any]] = []
    offset = args.offset
    while len(all_items) < args.limit:
        url = history_url_from_parts(args, offset, min(args.count, args.limit - len(all_items)))
        try:
            resp = request_with_retry(session, url, retries=args.retries)
            payload = resp.json()
        except Exception as exc:
            (output / "history-error.txt").write_text(str(exc), encoding="utf-8")
            print(f"ERROR: 历史列表请求失败: {exc}", file=sys.stderr)
            return 1
        if payload.get("ret") not in (0, "0", None):
            (output / "history-response.json").write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
            print(f"ERROR: 微信接口返回异常 ret={payload.get('ret')} errmsg={payload.get('errmsg')}", file=sys.stderr)
            return 1
        items, can_continue = parse_history_items(payload)
        if not items:
            break
        all_items.extend(items)
        offset += len(items)
        if not can_continue:
            break
        if args.delay_max > 0:
            time.sleep(random.uniform(args.delay_min, args.delay_max))
    all_items = all_items[: args.limit]
    links_path = output / "history-links.txt"
    links_path.write_text("\n".join(item["link"] for item in all_items) + "\n", encoding="utf-8")
    (output / "history-items.json").write_text(json.dumps(all_items, ensure_ascii=False, indent=2), encoding="utf-8")
    if args.archive and all_items:
        batch_args = argparse.Namespace(
            links=str(links_path),
            output=str(output / "archive"),
            resume=True,
            dry_run=args.dry_run,
            retries=args.retries,
            cookie=args.cookie,
            formats=args.formats,
            article_zip=args.article_zip,
            zip=args.zip,
            delay_min=args.delay_min,
            delay_max=args.delay_max,
            incremental=getattr(args, "incremental", False),
            catalog=getattr(args, "catalog", ""),
        )
        return batch_archive(batch_args)
    print(str(links_path))
    return 0


def current_archive(args: argparse.Namespace) -> int:
    if "mp.weixin.qq.com" not in args.url:
        print("ERROR: 仅支持 mp.weixin.qq.com 公开文章链接。", file=sys.stderr)
        return 2
    try:
        article_dir, _ = archive_article(
            args.url,
            Path(args.output),
            dry_run=args.dry_run,
            retries=args.retries,
            cookie=args.cookie or "",
            formats=set(args.formats.split(",")),
            make_zip=args.zip,
        )
    except Exception as exc:
        print(f"ERROR: {exc}", file=sys.stderr)
        return 1
    print(str(article_dir))
    return 0


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description="Archive public WeChat articles as Markdown/HTML and assets.")
    sub = parser.add_subparsers(dest="command")

    current = sub.add_parser("current", help="Archive one article link.")
    current.add_argument("url")
    current.add_argument("--output", default=".", help="Output directory; defaults to the current working directory so the article lands directly under ./<account>/<title>/.")
    current.add_argument("--formats", default="md,html,json", help="Comma separated: md,html,json,docx,pdf")
    current.add_argument("--cookie", default="", help="Optional user-provided legal Cookie.")
    current.add_argument("--retries", type=int, default=2)
    current.add_argument("--dry-run", action="store_true")
    current.add_argument("--zip", action="store_true", help="Zip the article directory.")
    current.set_defaults(func=current_archive)

    batch = sub.add_parser("batch", help="Archive explicit article links from a text file.")
    batch.add_argument("--links", required=True)
    batch.add_argument("--output", default="wechat-batch-archive")
    batch.add_argument("--formats", default="md,html,json")
    batch.add_argument("--cookie", default="")
    batch.add_argument("--retries", type=int, default=2)
    batch.add_argument("--resume", action="store_true")
    batch.add_argument("--dry-run", action="store_true")
    batch.add_argument("--article-zip", action="store_true")
    batch.add_argument("--zip", action="store_true", help="Zip the whole batch directory.")
    batch.add_argument("--delay-min", type=float, default=1.0)
    batch.add_argument("--delay-max", type=float, default=3.0)
    batch.add_argument("--incremental", action="store_true", help="Use catalog.json to skip articles already archived.")
    batch.add_argument("--catalog", default="", help="Optional catalog path for cross-run dedup/incremental updates.")
    batch.set_defaults(func=batch_archive)

    album = sub.add_parser("album", help="Extract article links from a public WeChat album page.")
    album.add_argument("url", help="WeChat album URL, e.g. https://mp.weixin.qq.com/mp/appmsgalbum?...")
    album.add_argument("--output", default="wechat-album")
    album.add_argument("--cookie", default="", help="Optional user-provided legal Cookie.")
    album.add_argument("--limit", type=int, default=0, help="Max links to export; 0 means no limit.")
    album.add_argument("--retries", type=int, default=2)
    album.add_argument("--archive", action="store_true", help="Archive exported links after discovery.")
    album.add_argument("--formats", default="md,html,json")
    album.add_argument("--dry-run", action="store_true")
    album.add_argument("--article-zip", action="store_true")
    album.add_argument("--zip", action="store_true", help="Zip the archive directory when --archive is used.")
    album.add_argument("--delay-min", type=float, default=1.0)
    album.add_argument("--delay-max", type=float, default=3.0)
    album.add_argument("--incremental", action="store_true", help="Use catalog.json to skip articles already archived when --archive is used.")
    album.add_argument("--catalog", default="", help="Optional catalog path for cross-run dedup/incremental updates.")
    album.set_defaults(func=fetch_album)

    workflow = sub.add_parser("workflow", help="Run album -> batch archive -> index/zip workflow.")
    workflow.add_argument("url", help="WeChat album URL.")
    workflow.add_argument("--output", default="wechat-workflow")
    workflow.add_argument("--cookie", default="")
    workflow.add_argument("--limit", type=int, default=20)
    workflow.add_argument("--formats", default="md,html,json")
    workflow.add_argument("--retries", type=int, default=2)
    workflow.add_argument("--dry-run", action="store_true")
    workflow.add_argument("--article-zip", action="store_true")
    workflow.add_argument("--zip", action="store_true")
    workflow.add_argument("--delay-min", type=float, default=1.0)
    workflow.add_argument("--delay-max", type=float, default=3.0)
    workflow.add_argument("--incremental", action="store_true", help="Use catalog.json to skip articles already archived when --archive is used.")
    workflow.add_argument("--catalog", default="", help="Optional catalog path for cross-run dedup/incremental updates.")
    workflow.set_defaults(func=workflow_archive)

    browser_context = sub.add_parser("browser-context", help="Extract links from a saved browser page export without reading cookies.")
    browser_context.add_argument("--input", required=True, help="JSON/HTML/text exported from browser_read_current_page or browser_read_page.")
    browser_context.add_argument("--output", default="wechat-browser-context")
    browser_context.add_argument("--limit", type=int, default=0)
    browser_context.add_argument("--archive", action="store_true")
    browser_context.add_argument("--formats", default="md,html,json")
    browser_context.add_argument("--retries", type=int, default=2)
    browser_context.add_argument("--dry-run", action="store_true")
    browser_context.add_argument("--article-zip", action="store_true")
    browser_context.add_argument("--zip", action="store_true")
    browser_context.add_argument("--delay-min", type=float, default=1.0)
    browser_context.add_argument("--delay-max", type=float, default=3.0)
    browser_context.add_argument("--incremental", action="store_true", help="Use catalog.json to skip articles already archived when --archive is used.")
    browser_context.add_argument("--catalog", default="", help="Optional catalog path for cross-run dedup/incremental updates.")
    browser_context.set_defaults(func=browser_context_extract)

    publish = sub.add_parser("publish", help="Create Lark/IMA-ready offline package from an archive directory.")
    publish.add_argument("--archive", required=True, help="Archive directory containing article.md files.")
    publish.add_argument("--target", choices=["lark", "ima", "both"], default="both")
    publish.add_argument("--output", default="wechat-publish-package")
    publish.add_argument("--mode", choices=["offline", "handoff", "online-request"], default="handoff", help="offline only packages files; handoff creates executable task manifests; online-request documents authorized write intent.")
    publish.add_argument("--destination", default="", help="Target Lark doc/wiki URL or IMA knowledge base name/id, if known.")
    publish.add_argument("--zip", action="store_true")
    publish.set_defaults(func=publish_archive)

    history = sub.add_parser("history", help="Fetch recent history links using user-provided login context.")
    history.add_argument("--profile-url", default="", help="Existing profile_ext URL copied from a legal browser session.")
    history.add_argument("--biz", default="")
    history.add_argument("--cookie", default="", help="Required legal Cookie from the user's own browser/session.")
    history.add_argument("--appmsg-token", default="")
    history.add_argument("--pass-ticket", default="")
    history.add_argument("--uin", default="")
    history.add_argument("--key", default="")
    history.add_argument("--scene", default="124")
    history.add_argument("--offset", type=int, default=0)
    history.add_argument("--count", type=int, default=10)
    history.add_argument("--limit", type=int, default=20)
    history.add_argument("--output", default="wechat-history")
    history.add_argument("--formats", default="md,html,json")
    history.add_argument("--retries", type=int, default=2)
    history.add_argument("--archive", action="store_true", help="Archive fetched links after discovery.")
    history.add_argument("--dry-run", action="store_true")
    history.add_argument("--article-zip", action="store_true")
    history.add_argument("--zip", action="store_true")
    history.add_argument("--delay-min", type=float, default=1.0)
    history.add_argument("--delay-max", type=float, default=3.0)
    history.add_argument("--incremental", action="store_true", help="Use catalog.json to skip articles already archived when --archive is used.")
    history.add_argument("--catalog", default="", help="Optional catalog path for cross-run dedup/incremental updates.")
    history.set_defaults(func=fetch_history)
    return parser


def main(argv: Optional[list[str]] = None) -> int:
    argv = list(argv if argv is not None else sys.argv[1:])
    if argv and argv[0].startswith("http"):
        argv = ["current"] + argv
    parser = build_parser()
    args = parser.parse_args(argv)
    if not hasattr(args, "func"):
        parser.print_help()
        return 2
    return args.func(args)


if __name__ == "__main__":
    raise SystemExit(main())
